from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import json

def create_list(paragraph, list_type):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), list_type) #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list
    pPr.append(numPr) #add number properties to paragraph

def loadJsonData(jsonFile):
    input_file = open(jsonFile)
    json_array = json.load(input_file)
    input_file.close()
    return json_array

def replace_in_line(paragraph,replace,content):
    toSkip = len(replace)
    pos = paragraph.text.index(replace)
    paragraph.text = paragraph.text[:pos]+content+paragraph.text[pos+toSkip:]


def check_new_section(currentOrder,possibleSection):
    try: 
        if currentOrder[1] == possibleSection:
            #Means we hit a new section
            currentOrder.pop(0)
    except IndexError:
        pass

    return currentOrder
    
def parse_json_basic(json,block,content):
    for i in json[block]:
        return i[content]

def order_of_blocks(document):
    order = []
    #Allows us to easily assess what section we are in 
    for paragraph in document.paragraphs:
        text = paragraph.text.lower()
        if text == "education":
            order.append("education")
        if text == "certifications":
            order.append("certifications")
        if text == "skills" or text == "skills & abilities":
            order.append("skills")
        if text == "links":
            order.append("links")
        if text == "experience":
            order.append("experience")
        if text == "projects":
            order.append("projects")
        if text == "end":
            order.append("end")
    return order

def skills(header,json):
    paragraph.text = ""
    paragraph.add_run(header).bold = True
    paragraph.add_run(": ")
    for lang in range(0,len(json)):
        try:
            test = json[lang+1]
            paragraph.add_run(json[lang] +", ")
        except IndexError:
            paragraph.add_run(json[lang])
                

json = loadJsonData('data/NoSqlSchema.json')
f = open('templates/template1.docx', 'rb')
#opens the document
document = Document(f)

usedDoubleMajor = False
usedAllCerts = False
usedMultipleProjects = False
usedExperience = False

styles = document.styles
paragraph_styles = [
    s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
 ]
style = styles['Heading 2']

        
#First Step,
order = order_of_blocks(document)

for paragraph in document.paragraphs:
    #rint(paragraph.text)
    text = paragraph.text.lower()
    order = check_new_section(order, text)
    #Name
    if 'firstname' in text: 
        paragraph.text = json["FirstName"]+ " " +json["LastName"]

    if 'address' in text:
        try:
            addressPos = text.index('address')
            #There is an address in the database
            address = json["Address"]
            paragraph.text = paragraph.text[:addressPos]+address+paragraph.text[addressPos+7:]
        except KeyError:
            #there is no address provided
            address = ""
            paragraph.text = paragraph.text[:addressPos]+address+paragraph.text[addressPos+10:]

    if 'phone' in text:
        replace_in_line(paragraph,"Phone",json['Phone']) 

    if 'email' in text:
        replace_in_line(paragraph,"Email",json['Email'])

    if 'summarytext' == text:
        paragraph.text = json["SummaryText"]

    if 'degree' in text:
        replace_in_line(paragraph,"Degree",parse_json_basic(json,"Education","Degree"))

    if 'graddate' in text and order[0] == "education":
        replace_in_line(paragraph,"GradDate",parse_json_basic(json,"Education","GradDate"))
    
    if 'school' in text and order[0] == "education":
        replace_in_line(paragraph,"School",parse_json_basic(json,"Education","School"))
    
    if 'major' in text and order[0] == "education":
        #just a check to see if they double majored in anything
        major = parse_json_basic(json,"Education","Major")
        paragraph.text += " "+major[0] 

    if 'minor' in text and order[0] == "education":
        minorList = parse_json_basic(json,"Education","Minor")
        for minors in range(len(minorList)):
            paragraph.text += minorList[minors] + ", "
        paragraph.text = paragraph.text[:-2]

    if "education" not in order and usedDoubleMajor == False:
        doubleMajor = parse_json_basic(json,"Education","Major")
        if len(doubleMajor) > 1:
            doubleDegree = True
        if doubleDegree == True:
            for degree in range(1,len(doubleMajor)):
                newDegree = paragraph.insert_paragraph_before(parse_json_basic(json,"Education","Degree")+" | "+parse_json_basic(json,"Education","GradDate")+" | " +parse_json_basic(json,"Education","School"))
                newDegree.style = style
                font = style.font
                font.name = 'Cambria'
                font.bold = True
                font.size = Pt(11)
                newMajor = paragraph.insert_paragraph_before("Major: "+doubleMajor[degree],style='List Bullet')
                create_list(newMajor, "1")
        usedDoubleMajor = True

    #Certificaation Block
    if 'certificationname' in text and order[0] == "certifications":  
        replace_in_line(paragraph,"CertificationName",parse_json_basic(json, "Certification","CertName"))

    if 'certificationassociation' in text and order[0] == "certifications":  
        replace_in_line(paragraph,"CertificationAssociation",parse_json_basic(json, "Certification","CertAssociation"))

    if 'certificationdate' in text and order[0] == "certifications":  
        replace_in_line(paragraph,"CertificationDate",parse_json_basic(json, "Certification","CertDate"))

    if "certifications" not in order and usedAllCerts == False:
        certificationList = json["Certification"]
        if len(certificationList) > 1:
            multipleCertifications = True
        if multipleCertifications == True:
            for cert in range(1,len(certificationList)):
                newCert = paragraph.insert_paragraph_before(json["Certification"][cert]["CertName"]+", "+json["Certification"][cert]["CertAssociation"]+" — " +json["Certification"][cert]["CertDate"])
                pstyle = document.styles['Normal']
                pfont = pstyle.font
                pfont.name = 'Cambria'
                pfont.size = Pt(11)
                pfont.color.rgb = RGBColor(64,64,64)
                create_list(newCert, "1")
                newCert.style = document.styles['Normal']
        usedAllCerts = True

    #Skills Block
    if "langskills" in text and "skills" == order[0]:
        skills("Languages", parse_json_basic(json, "Skill","Languages"))

    if "frameskills" in text and "skills" == order[0]:
        skills("Frameworks", parse_json_basic(json, "Skill","Frameworks"))

    if "techskills" in text and "skills" == order[0]:
        skills("Technologies & Software", parse_json_basic(json, "Skill","DatabaseTech"))

    #Project Block

    if 'projname' in text and order[0] == "projects":
        replace_in_line(paragraph,"ProjName",parse_json_basic(json,"Project","ProjName"))

    if 'projassociation' in text and order[0] == "projects":
        replace_in_line(paragraph,"ProjAssociation",parse_json_basic(json,"Project","ProjAssociation"))
    
    if 'projdate' in text and order[0] == "projects":
        replace_in_line(paragraph,"ProjDate",parse_json_basic(json,"Project","ProjDate"))
    
    if 'projinfo1' in text and order[0] == "projects":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"ProjInfo1",parse_json_basic(json,"Project","ProjInfo1"))

    if 'projinfo2' in text and order[0] == "projects":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"ProjInfo2",parse_json_basic(json,"Project","ProjInfo2"))

    if 'projinfo3' in text and order[0] == "projects":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"ProjInfo3",parse_json_basic(json,"Project","ProjInfo3"))

    if "projects" not in order and usedMultipleProjects == False:
        multiProj = json["Project"]
        if len(multiProj) > 1:
            moreProj = True
        if moreProj == True:
            for proj in range(1,len(multiProj)):
                newProj = paragraph.insert_paragraph_before(multiProj[proj]["ProjName"]+" | "+multiProj[proj]["ProjAssociation"]+" | " +multiProj[proj]["ProjDate"])
                newProj.style = style
                font = style.font
                font.name = 'Cambria'
                font.bold = True
                font.size = Pt(11)
                projInfo1 = paragraph.insert_paragraph_before(multiProj[proj]["ProjInfo1"],style='List Bullet')
                projInfo2 = paragraph.insert_paragraph_before(multiProj[proj]["ProjInfo2"],style='List Bullet')
                projInfo3 = paragraph.insert_paragraph_before(multiProj[proj]["ProjInfo3"],style='List Bullet')
                create_list(projInfo1, "1")
                create_list(projInfo2, "1")
                create_list(projInfo3, "1")
        usedMultipleProjects = True

    #Experience Block
    if 'jobtitle' in text and order[0] == "experience":
        replace_in_line(paragraph,"JobTitle",parse_json_basic(json,"Experience","JobTitle"))

    if 'joborg' in text and order[0] == "experience":
        replace_in_line(paragraph,"JobOrg",parse_json_basic(json,"Experience","JobOrg"))
    
    if 'startdate' in text and order[0] == "experience":
        replace_in_line(paragraph,"StartDate",parse_json_basic(json,"Experience","JobStartDate"))
    
    if 'enddate' in text and order[0] == "experience":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"EndDate",parse_json_basic(json,"Experience","JobEndDate"))

    if 'jobinfo1' in text and order[0] == "experience":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"JobInfo1",parse_json_basic(json,"Experience","JobInfo1"))

    if 'jobinfo2' in text and order[0] == "experience":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"JobInfo2",parse_json_basic(json,"Experience","JobInfo2"))

    if 'jobinfo3' in text and order[0] == "experience":
        #just a check to see if they double majored in anything
        replace_in_line(paragraph,"JobInfo3",parse_json_basic(json,"Experience","JobInfo3"))

    if "experience" not in order and usedExperience == False:
        multiExp = json["Experience"]
        if len(multiExp) > 1:
            moreExp = True
        if moreExp == True:
            for exp in range(1,len(multiExp)):
                newExp = paragraph.insert_paragraph_before(multiExp[exp]["JobTitle"]+" | "+multiExp[exp]["JobOrg"]+" | " +multiExp[exp]["JobStartDate"]+" - " +multiExp[exp]["JobEndDate"])
                newExp.style = style
                font = style.font
                font.name = 'Cambria'
                font.bold = True
                font.size = Pt(11)
                jobInfo1 = paragraph.insert_paragraph_before(multiExp[exp]["JobInfo1"],style='List Bullet')
                jobInfo2 = paragraph.insert_paragraph_before(multiExp[exp]["JobInfo2"],style='List Bullet')
                jobInfo3 = paragraph.insert_paragraph_before(multiExp[exp]["JobInfo3"],style='List Bullet')
                create_list(jobInfo1, "1")
                create_list(jobInfo2, "1")
                create_list(jobInfo3, "1")
        usedExperience = True

    if "end" == order[0]:
        paragraph.text = ""





document.save("demo.docx")
        


"""

    if "skills & abilities" in text:
        if len(degrees) > 1:
            for degree in range(1,len(degrees)):
                newDegree = paragraph.insert_paragraph_before(degrees[degree].strip())
                newDegree.style = style
                font = style.font
                font.name = 'Cambria'
                font.bold = True
                font.size = Pt(11)

                newMajor = paragraph.insert_paragraph_before("Major:",style='List Bullet')
                create_list(newMajor, "1")
                
                newMinor = paragraph.insert_paragraph_before("Minor:",style='List Bullet')
                create_list(newMinor,"1")




For columns
section = document.sections[0]

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

    #input()
    ## For Editing 
    #text = paragraph.text.lower()
    #if "project name" in text:
     #   print(paragraph.text)
     #   paragraph.text = "moo"


    #For removing stuff

    #if paragraph.text == "Project2":
     #   p = paragraph._element
      #  p.getparent().remove(p)
       # p._p = p._element = None
"""